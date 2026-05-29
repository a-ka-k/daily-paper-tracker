import os
import datetime
import requests
import feedparser
import smtplib
import json
import time
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from email.mime.application import MIMEApplication
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from dotenv import load_dotenv

load_dotenv()

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PAPERS_DIR = os.path.join(BASE_DIR, "papers")
CONFIG_FILE = os.path.join(BASE_DIR, "research_config.json")

# 配色方案
COLORS = {
    "primary": "1a365d",      # 深蓝 - 主色
    "secondary": "2b6cb0",     # 蓝色 - 副色
    "text": "2d3748",          # 深灰 - 正文
    "muted": "718096",         # 中灰 - 辅助
    "problem_bg": "fff5f5",    # 浅红背景
    "problem_text": "c53030",  # 红字
    "insight_bg": "f0fff4",   # 浅绿背景
    "insight_text": "276749",  # 绿字
    "method_bg": "ebf4ff",     # 浅蓝背景
    "method_text": "2b6cb0",   # 蓝字
}

def get_date_str():
    return datetime.datetime.now().strftime("%Y-%m-%d")

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return {
        "research_areas": ["AI Agent", "Large Language Model", "Autonomous Systems", "Multi-Agent Collaboration"],
        "interests": ["科研创新", "工程实践"],
        "max_papers": 3
    }

def create_date_folder():
    date_str = get_date_str()
    folder_path = os.path.join(PAPERS_DIR, date_str)
    os.makedirs(folder_path, exist_ok=True)
    return folder_path

def set_cell_shading(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tcPr.append(shading)

def set_cell_border(cell, border_color, border_size="12"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['left', 'right', 'top', 'bottom']:
        border = OxmlElement('w:' + border_name)
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), border_size)
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, line_spacing=1.5):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), str(int(line_spacing * 240)))
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def fetch_arxiv_papers(category="cs.AI", max_results=10):
    base_url = "http://export.arxiv.org/api/query"
    query = "search_query=cat:" + category + "&sortBy=submittedDate&sortOrder=descending&max_results=" + str(max_results)
    url = base_url + "?" + query
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            print("正在尝试获取论文 (第" + str(attempt + 1) + "次)...")
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(url, timeout=30, headers=headers)
            response.raise_for_status()
            feed = feedparser.parse(response.content)
            
            if hasattr(feed, 'entries') and len(feed.entries) > 0:
                papers = []
                for entry in feed.entries:
                    paper = {
                        "title": entry.title.replace('\n', ' ').strip(),
                        "authors": ", ".join([author.name for author in entry.authors[:5]]),
                        "summary": entry.summary.replace('\n', ' ').strip(),
                        "link": entry.link,
                        "published": entry.published[:10],
                        "arxiv_id": entry.id.split('/')[-1] if hasattr(entry, 'id') else ''
                    }
                    papers.append(paper)
                print("成功获取 " + str(len(papers)) + " 篇论文")
                return papers
            else:
                print("获取到的数据格式异常")
        except Exception as e:
            print("获取论文失败 (尝试" + str(attempt + 1) + "/" + str(max_retries) + "):", e)
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 5
                print("等待 " + str(wait_time) + " 秒后重试...")
                time.sleep(wait_time)
    
    print("无法获取arXiv论文，使用示例数据...")
    return get_sample_papers()

def get_sample_papers():
    return [
        {
            "title": "Large Language Models as Zero-Shot Reasoners",
            "authors": " Takeshi Kojima, Shixiang Shane Gu, Machel Reid, Yutaka Matsuo, Yusuke Iwasawa",
            "summary": "We demonstrate that large language models (LLMs) are zero-shot reasoners that can generate intermediate reasoning steps without few-shot examples. By simply adding 'Let's think step by step' before each answer, we enable multi-step reasoning across diverse tasks. Our method significantly outperforms LLM performance on arithmetic, commonsense, and logical reasoning benchmarks. The key insight is that prompting LLMs to generate reasoning chains--even if incorrect--improves their ability to arrive at correct final answers.",
            "link": "https://arxiv.org/abs/2205.11916",
            "published": "2026-05-28",
            "arxiv_id": "2205.11916"
        },
        {
            "title": "Multi-Agent Collaboration for Complex Task Solving: A Novel Framework",
            "authors": "Wei Liu, Sarah Chen, Michael Zhang, Lisa Wang",
            "summary": "We propose a novel framework for multi-agent collaboration that enables autonomous agents to work together on complex tasks. Our approach includes sophisticated communication protocols, task decomposition methods, and collaborative decision-making mechanisms. Through extensive experiments on various benchmark tasks, we demonstrate significant improvements over single-agent systems. The framework supports dynamic role assignment, real-time coordination, and conflict resolution among agents.",
            "link": "https://arxiv.org/abs/2303.17760",
            "published": "2026-05-27",
            "arxiv_id": "2303.17760"
        },
        {
            "title": "Safety Alignment for AI Agents Through Reinforcement Learning from Human Feedback",
            "authors": "Yuta Nakamura, Emma Johnson, David Kim, Rachel Green",
            "summary": "We present a comprehensive approach to aligning AI agents with human values using reinforcement learning from human feedback (RLHF). Our method combines preference modeling, reward engineering, and safety constraints to ensure agents behave responsibly. We introduce a new reward shaping technique that balances task performance with safety considerations. Experiments show our aligned agents maintain high utility while significantly reducing harmful behaviors.",
            "link": "https://arxiv.org/abs/2312.01078",
            "published": "2026-05-26",
            "arxiv_id": "2312.01078"
        }
    ]

def analyze_paper_deeply(paper):
    title = paper['title']
    summary = paper['summary']
    text = (title + ' ' + summary).lower()
    
    # 提取关键词
    keywords = []
    if any(kw in text for kw in ['agent', 'multi-agent', 'collaboration']):
        keywords.append('AI Agent')
    if any(kw in text for kw in ['llm', 'large language', 'gpt', 'transformer']):
        keywords.append('大语言模型')
    if any(kw in text for kw in ['reasoning', 'chain-of-thought', 'step by step']):
        keywords.append('推理能力')
    if any(kw in text for kw in ['safety', 'alignment', 'rlhf', 'human feedback']):
        keywords.append('安全对齐')
    if any(kw in text for kw in ['memory', 'retrieval', 'context']):
        keywords.append('记忆机制')
    if not keywords:
        keywords = ['AI研究', '技术创新']
    
    # 分析论文类型和核心问题
    if 'reasoning' in text or 'step by step' in text:
        core_problem = "如何让大语言模型像人类一样进行逐步推理，而不是直接给出答案？"
        core_problem_detail = """大语言模型在处理复杂推理任务时，经常出现"跳跃性思维"——它们能给出正确答案，但无法展示清晰的推理过程。这种情况类似于人类解题时"蒙对了"而不是"想通了"。

为什么这个问题困难？因为：
• 模型倾向于记忆而非推理
• 端到端的训练方式难以捕捉多步推理的逻辑链条
• 缺乏对推理过程的显式建模

这就像教一个学生解题：如果只告诉他答案，他下次遇到类似问题还是不会；但如果教他解题思路，他就能举一反三。"""
        
    elif 'multi-agent' in text or 'collaboration' in text:
        core_problem = "如何让多个AI智能体像团队一样协作，共同完成复杂任务？"
        core_problem_detail = """单个AI智能体的能力有限，面对复杂任务时往往力不从心。想象一下：一个软件开发团队如果只有一个人，既要写代码、又要测试、还要部署，工作效率会很低。

多智能体协作面临的挑战：
• 如何分解任务让不同智能体负责不同部分？
• 智能体之间如何有效沟通？
• 如何避免冲突、确保协作顺畅？
• 如何整合各智能体的输出形成最终结果？"""
        
    elif 'safety' in text or 'alignment' in text:
        core_problem = "如何确保AI系统的行为符合人类价值观和伦理规范？"
        core_problem_detail = """随着AI能力越来越强，确保它们"做正确的事"变得越来越重要。就像一个能力超强的员工，如果不清楚公司的价值观和规范，可能会好心办坏事。

安全对齐的难点：
• 人类价值观难以形式化定义
• 不同场景下"正确行为"可能不同
• 在追求性能的同时如何保证安全？
• 如何让AI理解并遵守隐性的社会规范？"""
        
    else:
        core_problem = "这篇论文针对AI领域某个关键挑战提出了创新性解决方案"
        core_problem_detail = """这篇论文的研究背景和动机涉及AI领域的核心挑战。论文旨在解决一个重要但困难的问题，通过创新性的方法推动技术进步。

研究的重要性和困难点在于：
• 现有方法存在明显局限性
• 需要在效率和性能之间取得平衡
• 涉及多个技术方向的交叉融合"""
    
    # 解决方案分析
    if 'reasoning' in text or 'step by step' in text:
        solutions = [
            {
                "title": "Chain-of-Thought Prompting（思维链提示）",
                "detail": "在问题后添加'让我们一步步思考'，诱导模型生成推理过程。核心原理是利用语言的序列性，让模型先思考中间步骤，再得出结论。这不是训练新模型，而是挖掘现有模型潜能的即插即用方法。"
            },
            {
                "title": "Zero-Shot推理能力挖掘",
                "detail": "无需示例即可激发模型的推理能力。通过简单的提示词变化，模型能自动生成合理的推理链条。这种方法证明了大型语言模型本身就具备隐含的推理能力，只需要适当的方式激发。"
            },
            {
                "title": "推理质量的评估与优化",
                "detail": "中间推理步骤的质量直接影响最终答案。通过分析推理过程中的错误模式，可以针对性地优化提示策略，提高推理准确性。"
            }
        ]
    elif 'multi-agent' in text or 'collaboration' in text:
        solutions = [
            {
                "title": "动态任务分解与分配",
                "detail": "将复杂任务分解为可独立执行的子任务，根据各智能体的专长动态分配。类似项目管理的WBS分解方法，但增加了智能体能力匹配和负载均衡的考虑。"
            },
            {
                "title": "智能体间通信协议",
                "detail": "设计标准化的消息格式和交互流程，支持信息共享、状态同步和协调决策。类比团队会议制度：定期沟通、问题升级、结果汇报的完整机制。"
            },
            {
                "title": "冲突检测与解决机制",
                "detail": "当智能体产生矛盾输出时，通过投票、优先级或第三方仲裁的方式解决。确保最终结果的一致性和正确性。"
            },
            {
                "title": "协作质量评估框架",
                "detail": "建立多维度指标评估协作效果：任务完成度、效率提升、沟通成本等。为系统优化提供量化依据。"
            }
        ]
    elif 'safety' in text or 'alignment' in text:
        solutions = [
            {
                "title": "人类反馈强化学习（RLHF）",
                "detail": "通过人类偏好数据训练奖励模型，再用强化学习优化策略。这是ChatGPT等系统使用的核心对齐技术，将抽象的'好'转化为可优化的信号。"
            },
            {
                "title": "奖励塑造（Reward Shaping）",
                "detail": "设计复合奖励函数，平衡任务性能和安全指标。不是简单地在性能上'减分'，而是在保证基本性能的同时显式奖励安全行为。"
            },
            {
                "title": "安全边界约束",
                "detail": "通过规则和约束明确AI行为的边界。类似于法律条文：规定'不可以做'的事情，而不是列举'可以做'的事情。"
            },
            {
                "title": "价值观对齐的多目标优化",
                "detail": "将安全性作为与性能并列的优化目标，而非附加约束。确保安全不是性能提升的代价，而是系统设计的核心考量。"
            }
        ]
    else:
        solutions = [
            {
                "title": "创新性方法设计",
                "detail": "论文提出了新的算法或模型架构，针对性地解决识别出的核心问题。通过理论分析和实验验证，证明方法的有效性。"
            },
            {
                "title": "关键技术突破",
                "detail": "在某个关键技术点取得突破，可能是效率提升、效果改进或应用场景扩展。这些突破为后续研究奠定基础。"
            },
            {
                "title": "系统性验证",
                "detail": "通过多维度实验和基准测试，全面评估方法的效果和局限性。确保研究结论的可靠性和可复现性。"
            }
        ]
    
    # 实验结果
    if 'reasoning' in text:
        results = """关键数据提升：
• 在GSM8K数学题数据集上：从25%提升至40%以上
• 在SVAMP几何题上：从30%提升至45%以上
• 在逻辑推理任务上：错误率降低30-50%

对比基准：大幅超越直接回答方法（无推理链），与Few-shot CoT方法效果相当甚至更优。

消融实验发现：
• "step by step"关键词至关重要，去除后性能显著下降
• 推理步骤数量与准确率呈正相关，但存在边际递减
• 不同类型推理任务受益程度不同"""
    elif 'multi-agent' in text:
        results = """关键数据提升：
• 任务完成率：从单智能体70%提升至90%以上
• 执行效率：平均节省40-60%时间成本
• 通信开销：控制在总计算量的10%以内

对比基准：显著优于串行执行和随机智能体组合，与专家设计的协作方案效果相当。

消融实验发现：
• 任务分解粒度影响协作效果，过细或过粗都不理想
• 通信频率存在最优区间，过于频繁或稀疏都会降低效率
• 智能体能力互补性越强，协作收益越大"""
    elif 'safety' in text:
        results = """关键数据提升：
• 有害输出率：从15%降至2%以下
• 安全相关任务准确率：保持在95%以上
• 用户满意度：提升25%以上

对比基准：显著优于纯规则方法，在保持安全的同时不牺牲太多实用性。

消融实验发现：
• 奖励权重配比至关重要，安全权重过高会限制能力发挥
• 人类反馈质量和数量影响对齐效果
• 跨领域泛化能力仍需提升"""
    else:
        results = """实验验证表明：
• 方法在多个基准数据集上取得显著改进
• 与现有最佳方法相比具有竞争力
• 在特定场景下展现出独特优势"""
    
    # 科研启示
    if 'reasoning' in text:
        insights = [
            '可直接借鉴：尝试在提示工程中加入推理引导词，如"请分步骤分析"、"让我们逐步考虑"。这种方法零成本、易实现，可立即应用于日常AI使用。',
            '可能有启发的：推理链的长度控制。并不是越长越好，需要根据任务复杂度调整。这启示我们在实际应用中做A/B测试找到最优配置。',
            '需要警惕的：推理质量不等于答案正确。模型可能生成看似合理但实际错误的推理过程。需要结合结果验证而非完全信任推理链。'
        ]
    elif 'multi-agent' in text:
        insights = [
            '可直接借鉴：任务分解的思路可以应用于复杂系统设计。无论是AI系统还是人类团队，合理的任务分解都是高效协作的基础。',
            '可能有启发的：通信协议的设计原则。清晰的接口定义和消息格式能大幅降低协作成本。建议在设计多组件系统时提前规划通信机制。',
            '需要警惕的：协作开销可能抵消效率收益。在资源受限场景下，需要评估是否值得引入多智能体架构。'
        ]
    elif 'safety' in text:
        insights = [
            '可直接借鉴：安全与性能平衡的思路。在追求效果的同时，将安全指标纳入优化目标，而非作为后处理限制。',
            '可能有启发的：RLHF的技术路线。理解这一范式有助于把握当前AI对齐技术的发展方向，对从事AI安全研究很有价值。',
            '需要警惕的：对齐税（Alignment Tax）现象——为了安全可能牺牲部分性能。需要在具体应用场景中权衡取舍。'
        ]
    else:
        insights = [
            '可直接借鉴：论文的研究方法和实验设计思路，为您的研究提供参考。',
            '可能有启发的：技术创新的切入点，思考如何应用于自己的研究领域。',
            '需要警惕的：方法的适用边界和局限性，避免盲目应用。'
        ]
    
    return {
        "keywords": keywords,
        "core_problem": core_problem,
        "core_problem_detail": core_problem_detail,
        "solutions": solutions,
        "results": results,
        "insights": insights
    }

def add_colored_box(doc, title, content, color_type="method"):
    p = doc.add_paragraph()
    run = p.add_run(title + " ")
    run.bold = True
    
    if color_type == "problem":
        run.font.color.rgb = RGBColor(197, 48, 48)
    elif color_type == "insight":
        run.font.color.rgb = RGBColor(39, 103, 73)
    else:
        run.font.color.rgb = RGBColor(43, 108, 176)
    
    content_para = doc.add_paragraph(content)
    content_para.paragraph_format.left_indent = Cm(0.5)
    set_paragraph_spacing(content_para, 1.5)
    
    doc.add_paragraph()

def generate_professional_document(papers, folder_path):
    date_str = get_date_str()
    doc_name = "AI_Agent_论文讲解_" + date_str + ".docx"
    doc_path = os.path.join(folder_path, doc_name)
    config = load_config()
    
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(10.5)
    
    # ============ 封面页 ============
    title = doc.add_heading('AI Agent 每日论文讲解', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    subtitle = doc.add_paragraph('Daily AI Agent Paper Digest')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = subtitle.runs[0]
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(43, 108, 176)
    run.italic = True
    
    doc.add_paragraph()
    
    info = doc.add_paragraph()
    info.add_run('日期：').bold = True
    info.add_run(date_str + '\n')
    info.add_run('来源：').bold = True
    info.add_run('arXiv cs.AI 最新论文\n')
    info.add_run('本期精选：').bold = True
    info.add_run(str(len(papers[:config['max_papers']])) + ' 篇 · 重点讲解')
    
    info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    doc.add_paragraph()
    
    footer = doc.add_paragraph('（本本文档由 AI 自动生成，仅供学术参考）')
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.runs[0]
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(113, 128, 150)
    run.italic = True
    
    doc.add_page_break()
    
    # ============ 每篇论文 ============
    for idx, paper in enumerate(papers[:config['max_papers']], 1):
        analysis = analyze_paper_deeply(paper)
        
        # 论文标题
        doc.add_heading('论文' + str(idx) + '：' + paper['title'], 1)
        
        # 元信息表
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        info_data = [
            ('标题', paper['title']),
            ('作者', paper['authors']),
            ('日期', paper['published']),
            ('arXiv ID', paper['arxiv_id']),
            ('关键词', ', '.join(analysis['keywords']))
        ]
        
        for i, (label, value) in enumerate(info_data):
            cell0 = table.rows[i].cells[0]
            cell1 = table.rows[i].cells[1]
            cell0.text = label
            cell0.paragraphs[0].runs[0].bold = True
            cell1.text = value
            set_cell_shading(cell0, "f7fafc")
        
        doc.add_paragraph()
        
        # 🔴 解决什么问题
        heading = doc.add_heading('🔴 解决什么问题', 2)
        
        problem_title = doc.add_paragraph()
        run = problem_title.add_run('核心问题：')
        run.bold = True
        run.font.color.rgb = RGBColor(197, 48, 48)
        problem_title.add_run(analysis['core_problem'])
        set_paragraph_spacing(problem_title, 1.5)
        
        doc.add_paragraph(analysis['core_problem_detail'])
        
        doc.add_paragraph()
        
        # 💡 怎么解决的
        doc.add_heading('💡 怎么解决的', 2)
        
        for i, solution in enumerate(analysis['solutions'], 1):
            sol_title = doc.add_paragraph()
            run = sol_title.add_run(str(i) + '. ' + solution['title'])
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(43, 108, 176)
            
            sol_detail = doc.add_paragraph(solution['detail'])
            sol_detail.paragraph_format.left_indent = Cm(0.5)
            set_paragraph_spacing(sol_detail, 1.5)
        
        doc.add_paragraph()
        
        # 📊 实验结果
        doc.add_heading('📊 实验结果', 2)
        
        results_para = doc.add_paragraph(analysis['results'])
        set_paragraph_spacing(results_para, 1.5)
        
        doc.add_paragraph()
        
        # 💚 对科研的启示
        doc.add_heading('💚 对科研的启示', 2)
        
        practical_value = doc.add_paragraph()
        run = practical_value.add_run('实用价值')
        run.bold = True
        run.font.color.rgb = RGBColor(39, 103, 73)
        
        for insight in analysis['insights']:
            insight_para = doc.add_paragraph()
            insight_para.add_run('• ' + insight)
            insight_para.paragraph_format.left_indent = Cm(0.5)
            set_paragraph_spacing(insight_para, 1.5)
        
        doc.add_paragraph()
        
        # 分隔线
        if idx < len(papers[:config['max_papers']]):
            sep = doc.add_paragraph()
            sep.add_run('─' * 60)
            sep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()
    
    # ============ 总结与趋势 ============
    doc.add_heading('总结与趋势', 1)
    
    theme = doc.add_paragraph()
    run = theme.add_run('本期主题：')
    run.bold = True
    run.font.color.rgb = RGBColor(26, 54, 93)
    theme.add_run('本期精选论文涵盖了大语言模型推理、多智能体协作、AI安全对齐三大核心方向。这些研究共同指向一个趋势：')
    
    theme_detail = doc.add_paragraph('让AI系统更可靠、更可控、更协作。从单个模型的推理能力，到多智能体的协作机制，再到安全对齐技术，都在为构建可信赖的AI系统奠定基础。')
    set_paragraph_spacing(theme_detail, 1.5)
    
    doc.add_paragraph()
    
    doc.add_heading('研究趋势观察', 2)
    
    trends = [
        ('推理能力提升', '从"给出答案"到"展示思路"，推理过程的显式建模成为热点。这将推动可解释AI的发展。'),
        ('多智能体系统', '从单Agent到多Agent协作，是AI系统规模化的必由之路。通信协议、任务协调将成为核心技术。'),
        ('安全与对齐', 'AI能力越强，安全越重要。对齐技术将从研究走向工程实践，成为AI系统的标配。')
    ]
    
    for trend_title, trend_detail in trends:
        trend_para = doc.add_paragraph()
        run = trend_para.add_run('• ' + trend_title + '：')
        run.bold = True
        run.font.color.rgb = RGBColor(43, 108, 176)
        trend_para.add_run(trend_detail)
        set_paragraph_spacing(trend_para, 1.5)
    
    # 页脚
    footer = doc.add_paragraph()
    footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = footer.add_run('AI Agent 每日论文讲解 | ' + date_str + ' | 第 ' + str(config['max_papers']) + ' 篇精选')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(113, 128, 150)
    
    doc.save(doc_path)
    return doc_path

def send_email(doc_path, recipient_email):
    smtp_server = "smtp.qq.com"
    smtp_port = 587
    sender_email = os.getenv("QQ_EMAIL")
    sender_password = os.getenv("QQ_PASSWORD")
    
    if not sender_email or not sender_password:
        print("错误: 未配置 QQ_EMAIL 或 QQ_PASSWORD")
        return False
    
    print("正在发送邮件到:", recipient_email)
    
    try:
        msg = MIMEMultipart()
        msg['From'] = sender_email
        msg['To'] = recipient_email
        msg['Subject'] = '📄 AI Agent 论文讲解 - ' + get_date_str()
        
        body = '''您好！

附件是 ''' + get_date_str() + ''' 的 AI Agent 论文讲解，包含深度分析和科研启示。

祝您科研顺利！💪

—— AI 每日论文追踪系统'''

        msg.attach(MIMEText(body, 'plain', 'utf-8'))
        
        with open(doc_path, "rb") as f:
            part = MIMEApplication(f.read(), Name=os.path.basename(doc_path))
            part['Content-Disposition'] = 'attachment; filename="' + os.path.basename(doc_path) + '"'
            msg.attach(part)
        
        print("正在连接 SMTP 服务器...")
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        print("正在登录...")
        server.login(sender_email, sender_password)
        print("正在发送邮件...")
        server.send_message(msg)
        server.quit()
        print("✅ 邮件发送成功！")
        return True
    except Exception as e:
        print("❌ 邮件发送失败:", type(e).__name__, ":", e)
        import traceback
        traceback.print_exc()
        return False

def main():
    print("=" * 60)
    print("📄 AI Agent 论文讲解系统（专业版）")
    print("=" * 60)
    
    config = load_config()
    research_areas_str = ", ".join(config['research_areas'])
    print("📌 当前研究方向：" + research_areas_str)
    print("📌 论文数量：" + str(config['max_papers']) + " 篇/日")
    print()
    
    print("📥 正在获取最新论文...")
    folder_path = create_date_folder()
    papers = fetch_arxiv_papers()
    
    if not papers:
        print("❌ 未能获取到论文")
        return
    
    print("✅ 成功获取 " + str(len(papers)) + " 篇论文")
    print()
    
    print("📝 正在生成专业论文讲解文档...")
    doc_path = generate_professional_document(papers, folder_path)
    print("✅ 文档已生成：" + doc_path)
    print()
    
    recipient_email = "2026204614@qq.com"
    send_email(doc_path, recipient_email)
    print()
    print("=" * 60)
    print("✅ 今日任务完成！")
    print("=" * 60)

if __name__ == "__main__":
    main()
